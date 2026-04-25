"""Dokumentum betoltes: PDF, DOCX, kep (PNG/JPG).

Strategiak:
- PDF: PyMuPDF natv szoveg. Ha tul keves szoveg jon vissza (< 50 karakter),
  Tesseract OCR-t hasznalunk (valoszinuleg szkennelt PDF).
- DOCX: python-docx natv bekezdes-iteracio.
- Kep: Tesseract OCR kozvetlenul (hun+eng+deu nyelvekkel).

A vissza adott Document dataclass egyseges: file_name + full_text + pages.
A pages lista segit a forras-hivatkozasban (page_number).
"""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Union

import fitz  # PyMuPDF

# pytesseract es Pillow csak OCR fallback-hez kell -- lusta import,
# hogy a modul barhol importalhato legyen (pl. regex-extract teszthez).

OCR_LANGUAGES = "hun+eng+deu"
MIN_TEXT_CHARS_PER_PAGE = 50


def _ocr_image_bytes(image_bytes: bytes) -> str:
    """OCR lusta importtal -- ha pytesseract/Pillow nincs, ures stringet ad."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""
    image = Image.open(BytesIO(image_bytes))
    try:
        return pytesseract.image_to_string(image, lang=OCR_LANGUAGES)
    except pytesseract.TesseractError:
        try:
            return pytesseract.image_to_string(image)
        except Exception:
            return ""
    except Exception:
        return ""


@dataclass
class Page:
    page_number: int  # 1-indexelt
    text: str


@dataclass
class Document:
    file_name: str
    full_text: str
    pages: list[Page] = field(default_factory=list)


def _ocr_page_image(pix: fitz.Pixmap) -> str:
    """PyMuPDF Pixmap -> bytes -> OCR (lusta import)."""
    return _ocr_image_bytes(pix.tobytes(output="png"))


def load_pdf(file_path: Union[str, Path], file_name: str | None = None) -> Document:
    """PDF betoltes PyMuPDF + Tesseract fallback."""
    path = Path(file_path)
    doc_name = file_name or path.name
    pages: list[Page] = []

    with fitz.open(path) as pdf:
        for i, page in enumerate(pdf, start=1):
            text = page.get_text() or ""
            if len(text.strip()) < MIN_TEXT_CHARS_PER_PAGE:
                # OCR fallback: szkennelt PDF vagy kep-alapu
                try:
                    pix = page.get_pixmap(dpi=200)
                    text = _ocr_page_image(pix) or text
                except Exception:
                    pass  # OCR failure -> maradunk az eredeti (esetleg ures) szoveggel
            pages.append(Page(page_number=i, text=text))

    full_text = "\n\n".join(p.text for p in pages).strip()
    return Document(file_name=doc_name, full_text=full_text, pages=pages)


def load_docx(file_path: Union[str, Path], file_name: str | None = None) -> Document:
    """DOCX betoltes python-docx-szel (lusta import)."""
    from docx import Document as DocxDocument

    path = Path(file_path)
    doc_name = file_name or path.name

    docx = DocxDocument(str(path))
    paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    # DOCX-nek nincs termesztett oldal-fogalma -> egyetlen "oldal"
    pages = [Page(page_number=1, text=full_text)]
    return Document(file_name=doc_name, full_text=full_text, pages=pages)


def load_image(file_path: Union[str, Path], file_name: str | None = None) -> Document:
    """Kepfajl (PNG/JPG) OCR-rel (lusta import)."""
    path = Path(file_path)
    doc_name = file_name or path.name
    with open(path, "rb") as fh:
        text = _ocr_image_bytes(fh.read())
    pages = [Page(page_number=1, text=text)]
    return Document(file_name=doc_name, full_text=text.strip(), pages=pages)


def load_bytes(file_name: str, content: bytes) -> Document:
    """Egy feltoltott fajlhoz -- Streamlit UploadedFile.getvalue() jon ide."""
    suffix = Path(file_name).suffix.lower()

    if suffix == ".pdf":
        with fitz.open(stream=content, filetype="pdf") as pdf:
            pages: list[Page] = []
            for i, page in enumerate(pdf, start=1):
                text = page.get_text() or ""
                if len(text.strip()) < MIN_TEXT_CHARS_PER_PAGE:
                    try:
                        pix = page.get_pixmap(dpi=200)
                        text = _ocr_page_image(pix) or text
                    except Exception:
                        pass
                pages.append(Page(page_number=i, text=text))
        full_text = "\n\n".join(p.text for p in pages).strip()
        return Document(file_name=file_name, full_text=full_text, pages=pages)

    if suffix == ".docx":
        # python-docx ZIP-et vár -> BytesIO
        from docx import Document as DocxDocument
        docx = DocxDocument(BytesIO(content))
        paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return Document(
            file_name=file_name, full_text=full_text,
            pages=[Page(page_number=1, text=full_text)],
        )

    if suffix in {".png", ".jpg", ".jpeg"}:
        text = _ocr_image_bytes(content)
        return Document(
            file_name=file_name, full_text=text.strip(),
            pages=[Page(page_number=1, text=text)],
        )

    if suffix == ".txt":
        text = content.decode("utf-8", errors="replace")
        return Document(
            file_name=file_name, full_text=text,
            pages=[Page(page_number=1, text=text)],
        )

    raise ValueError(f"Nem tamogatott fajlkiterjesztes: {suffix}")
